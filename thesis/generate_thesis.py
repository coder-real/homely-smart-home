#!/usr/bin/env python3
"""Convert thesis to Word - clean rebuild."""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

BASE = '/home/user/homely-thesis'
OUT = os.path.join(BASE, 'thesis.docx')

doc = Document()

# Page setup
for sec in doc.sections:
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2.54)

# Default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0

# Helper functions
def txt(text, bold=False, italic=False, size=12, align=None, after=0, before=0, indent=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if align: p.alignment = align
    if after: p.paragraph_format.space_after = Pt(after)
    if before: p.paragraph_format.space_before = Pt(before)
    if indent: p.paragraph_format.first_line_indent = Cm(indent)
    return p

def fig(path, caption):
    full = os.path.join(BASE, path)
    if not os.path.exists(full):
        txt(f'[Missing: {path}]', align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(full, width=Inches(5))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run(caption)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.bold = True
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def eq(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    txt(text, indent=1.27, after=12)

def pagebreak():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()

txt('Design and Implementation of an ESP32-Based IoT Home Automation System with Mobile Application Control',
    bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
txt('Thesis submitted for the degree of Doctor of Philosophy',
    size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
txt('Joseph Sarwuan Tarka University',
    size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
txt('Department of Electrical and Electronics Engineering',
    size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
txt('By:', size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
txt('[Student Full Name]', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
txt('[Matriculation Number]', size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
txt('Supervisor:', size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
txt("[Supervisor's Full Name]", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
txt('[Month, Year]', size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
pagebreak()

# ══════════════════════════════════════════════════════════════
# DECLARATION
# ══════════════════════════════════════════════════════════════
txt('DECLARATION', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
body('I hereby declare that this thesis is my original work and has not been submitted for any other degree or diploma at this or any other university. All sources of information used have been acknowledged.')
txt('')
txt('[Student Full Name]', after=6)
txt('[Date]', after=24)
pagebreak()

# ══════════════════════════════════════════════════════════════
# CERTIFICATION
# ══════════════════════════════════════════════════════════════
txt('CERTIFICATION', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
body('This thesis has been approved for submission to the Department of Electrical and Electronics Engineering, Joseph Sarwuan Tarka University, Makurdi.')
txt('________________________                    ________________', after=6)
txt('Supervisor                                        Date', after=24)
txt('')
txt('________________________                    ________________', after=6)
txt('Head of Department                           Date', after=24)
pagebreak()

# ══════════════════════════════════════════════════════════════
# DEDICATION
# ══════════════════════════════════════════════════════════════
txt('DEDICATION', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
txt('To my parents, for their unwavering support and patience throughout my academic journey.',
    align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
pagebreak()

# ══════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ══════════════════════════════════════════════════════════════
txt('ACKNOWLEDGEMENTS', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for t in [
    "I would like to express my sincere gratitude to my supervisor, [Supervisor's Name], for the guidance, patience, and encouragement throughout this research.",
    "I am grateful to the Department of Electrical and Electronics Engineering at Joseph Sarwuan Tarka University for providing the laboratory facilities and technical resources.",
    "I acknowledge my colleagues in the research group for the many informal discussions that helped refine my thinking on IoT system design.",
    "I thank my family for their unwavering support and patience during the long hours spent in the laboratory and at the writing desk.",
    "Finally, I acknowledge the broader open-source community whose work on ESP32 libraries, React Native, and the Arduino ecosystem made this project feasible."
]:
    body(t)
pagebreak()

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════
txt('ABSTRACT', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
body(
    "The proliferation of Internet of Things (IoT) technologies has created opportunities for intelligent home automation systems "
    "that enhance comfort, energy efficiency, and security. However, existing commercial solutions often rely on cloud-based "
    "architectures that introduce latency, privacy concerns, and single points of failure. This thesis presents the design, "
    "implementation, and evaluation of a locally-controlled IoT home automation system built around the ESP32 microcontroller "
    "platform. The system integrates motion-triggered lighting via passive infrared (PIR) sensors, environmental monitoring "
    "through temperature and humidity sensing, and relay-based actuator control for multiple domestic devices. A custom mobile "
    "application developed using React Native and Expo provides real-time monitoring and control through a local REST API, "
    "eliminating cloud dependency. The firmware implements a captive portal for zero-configuration Wi-Fi provisioning, mDNS "
    "and UDP-based device discovery, and a dual-mode control architecture. Experimental validation demonstrates sub-second "
    "response times for relay actuation, 97% motion detection accuracy, and 99.6% state synchronization accuracy. The total "
    "system cost remains below N22,500, making it accessible for deployment in resource-constrained environments."
)
pagebreak()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
txt('TABLE OF CONTENTS', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for item, sub in [
    ("Declaration", False), ("Certification", False), ("Dedication", False),
    ("Acknowledgements", False), ("Abstract", False), ("List of Figures", False),
    ("List of Tables", False), ("Chapter 1: Introduction", False),
    ("1.1. Overview", True), ("1.2. Internet of Things", True),
    ("1.3. Motivation and Problem Statement", True), ("1.4. Research Objectives", True),
    ("1.5. Research Contributions", True), ("1.6. Thesis Outline", True),
    ("Chapter 2: Literature Review", False), ("Chapter 3: System Architecture and Design", False),
    ("3.7. Component Selection Calculations", True),
    ("Chapter 4: Implementation", False), ("Chapter 5: Results and Evaluation", False),
    ("Chapter 6: Concluding Remarks", False), ("Chapter 7: Appendix", False),
    ("Chapter 8: References", False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if sub: p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(item)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    if not sub: r.bold = True
pagebreak()

# ══════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ══════════════════════════════════════════════════════════════
txt('LIST OF FIGURES', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for f, t in [
    ("1-1", "Global smart home market growth"), ("1-2", "IoT architecture layers"),
    ("2-1", "Home automation evolution timeline"), ("2-3", "Microcontroller comparison"),
    ("2-4", "PIR detection zones"), ("2-6", "mDNS discovery sequence"),
    ("2-7", "UDP broadcast discovery"), ("2-8", "Cloud vs local architecture"),
    ("3-1", "System architecture"), ("3-2", "ESP32 wiring"), ("3-3", "Relay wiring"),
    ("3-5", "Power supply"), ("3-8", "State machine"), ("3-10", "Polling sequence"),
    ("3-11", "Optimistic UI flow"), ("3-12", "Communication flow"),
    ("4-3", "Voltage vs current"), ("4-4", "Wi-Fi provisioning"), ("4-6", "Temperature hysteresis"),
    ("5-2", "PIR accuracy"), ("5-3", "DHT11 accuracy"), ("5-4", "Power consumption"),
    ("5-5", "HTTP latency"), ("5-6", "Wi-Fi stability"), ("5-7", "Discovery time"),
    ("5-8", "End-to-end latency"), ("5-9", "Sync accuracy"), ("5-10", "7-day operation"),
    ("5-11", "Cost comparison"), ("5-12", "Radar comparison"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f'Figure {f}\t{t}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
pagebreak()

# ══════════════════════════════════════════════════════════════
# LIST OF TABLES
# ══════════════════════════════════════════════════════════════
txt('LIST OF TABLES', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
for t, n in [
    ("3-1", "ESP32 GPIO pin assignments"), ("3-4", "LED resistor calculations"),
    ("3-5", "System current budget"), ("4-1", "Power supply measurements"),
    ("5-1", "HTTP latency"), ("5-2", "Bill of materials"), ("5-3", "Comparison"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(f'Table {t}\t{n}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
pagebreak()

# ══════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════
h1('1. Introduction')

h2('1.1. Overview')
body('The concept of home automation has evolved significantly over the past two decades. The emergence of the Internet of Things (IoT) paradigm has been the primary catalyst for this transformation, enabling everyday household devices to connect to networks, exchange data, and respond to user commands remotely [1]. The global smart home market was valued at approximately $79 billion in 2020 and is projected to reach $313 billion by 2026 [2].')
fig('figures/figure-1-1-market-growth.png', 'Figure 1-1: Global smart home market value growth from 2018 to 2026')
body('The ESP32 microcontroller has emerged as a popular platform for IoT applications due to its integrated Wi-Fi and Bluetooth capabilities, dual-core processor, and low cost [6]. This thesis presents the design, implementation, and evaluation of a complete IoT home automation system built around the ESP32 platform.')

h2('1.2. Internet of Things and Smart Home Technologies')
body('The Internet of Things refers to the network of physical objects embedded with sensors, software, and connectivity capabilities [8]. The architecture can be categorized into three layers: perception (sensors and actuators), network (communication protocols), and application (user interfaces) [9].')
fig('figures/figure-1-2-iot-architecture.png', 'Figure 1-2: Three-layer IoT architecture')

h2('1.3. Motivation and Problem Statement')
body('Cloud dependency. Most commercial platforms rely on cloud infrastructure, introducing latency, privacy concerns, and reliability issues [17].')
body('Device discovery. Connecting IoT devices typically requires manual configuration of Wi-Fi credentials and IP addresses [19].')
body('User interface responsiveness. Cloud-mediated communication introduces 200-500ms latency [20].')
body('Cost and accessibility. Commercial solutions require proprietary hubs and subscriptions [21].')
body('Power supply reliability. In Nigeria, frequent power outages make internet-dependent systems unreliable. A locally-controlled system provides continuity independent of internet availability.')

h2('1.4. Research Objectives')
for i, o in enumerate([
    'To design and implement an ESP32-based home automation system operating entirely on the local network without cloud dependency.',
    'To develop a mobile application with automatic device discovery and responsive real-time control through optimistic UI updates.',
    'To evaluate the system\'s performance in terms of response latency, detection accuracy, and connection reliability.',
], 1):
    txt(f'{i}. {o}', indent=1.27, after=6)

h2('1.5. Research Contributions')
for i, (title, desc) in enumerate([
    ('A fully locally-controlled home automation architecture', 'that eliminates cloud dependency.'),
    ('A hybrid device discovery mechanism', 'combining mDNS and UDP broadcast fallback.'),
    ('An optimistic UI update pattern with server-side rollback', 'for responsive IoT control.'),
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'{i}. {title} ')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    r = p.add_run(desc)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

pagebreak()

# ══════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════
h1('3. System Architecture and Design')

h2('3.1. System Overview')
body('The system consists of three components: the ESP32 hardware controller, the firmware, and the mobile application.')
fig('figures/figure-3-1-system-architecture.png', 'Figure 3-1: System architecture block diagram')

h2('3.2. Hardware Architecture')
h3('3.2.1. ESP32 Configuration')
body('The GPIO pin assignments are presented in Table 3-1.')
tbl(['GPIO', 'Function', 'Notes'], [
    ['4', 'PIR motion sensor', 'Digital input'],
    ['17', 'DHT11 data', 'Single-wire protocol'],
    ['25', 'Relay - Bedroom fan', 'Active-LOW'],
    ['26', 'Relay - Porch light', 'Active-LOW'],
    ['32', 'Relay - Living room', 'Active-LOW'],
    ['33', 'Relay - Bedroom light', 'Active-LOW'],
])
txt('Table 3-1: ESP32 GPIO pin assignments', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
fig('figures/figure-3-2-esp32-wiring.png', 'Figure 3-2: ESP32 GPIO pin connections')
fig('figures/figure-3-3-relay-wiring.png', 'Figure 3-3: Relay module wiring schematic')
fig('figures/figure-3-5-power-supply-circuit.png', 'Figure 3-5: Power supply circuit schematic')

h2('3.7. Component Selection and Design Calculations')

h3('3.7.1. LED Current-Limiting Resistor Calculation')
body('LEDs require a current-limiting resistor. The value is determined by Ohm\'s law:')
eq('R = (V_supply - V_f) / I_desired')
body('Porch 1W LED (12V supply): V_supply = 12V, V_f = 3.2V, I_desired = 40mA')
eq('R = (12 - 3.2) / 0.040 = 220 ohms')
eq('P_R = I^2 x R = (0.040)^2 x 220 = 0.352W')
body('A 3W-rated resistor is selected (8.5x safety margin).')
body('Room 10mm LEDs (12V supply): V_supply = 12V, V_f = 3.0V, I_desired = 19mA')
eq('R = (12 - 3.0) / 0.019 = 473.7 ohms (nearest standard: 470 ohms)')
eq('I = (12 - 3.0) / 470 = 19.1mA')
eq('P_R = (0.0191)^2 x 470 = 0.172W')

tbl(['LED', 'Supply (V)', 'V_f (V)', 'I (mA)', 'R (ohms)', 'P_R (W)', 'Rating (W)'], [
    ['Porch 1W', '12', '3.2', '40', '220', '0.35', '3'],
    ['Room 10mm', '12', '3.0', '19', '470', '0.17', '1'],
])
txt('Table 3-4: LED resistor calculations', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

h3('3.7.2. Relay Coil Current')
eq('I_coil = V / R = 5V / 70 ohms = 71.4mA per coil')
eq('I_total = 4 x 71.4 = 285.6mA')

tbl(['Component', 'Current (mA)'], [
    ['ESP32 (Wi-Fi on)', '180'],
    ['Relay coils (4x)', '286'],
    ['Sensors + LEDs', '60'],
    ['Total', '526'],
])
txt('Table 3-5: System current budget', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

h3('3.7.3. 7805 Thermal Analysis')
eq('P = (V_in - V_out) x I = (12 - 5) x 0.526 = 3.68W')
eq('Without heatsink: Delta T = 3.68 x 50 = 184 deg C (exceeds limit)')
eq('With heatsink: Delta T = 3.68 x 20 = 73.6 deg C')
eq('T_junction = 25 + 73.6 = 98.6 deg C (within 125 deg C limit)')

h3('3.7.4. PIR Debounce Timing')
eq('t_debounce = N x T = 3 x 50ms = 150ms')

h3('3.7.5. Fan Hysteresis')
eq('Hysteresis = T_ON - T_OFF = 28.0 - 26.5 = 1.5 deg C')
fig('figures/figure-4-6-temperature-hysteresis.png', 'Figure 4-6: Temperature hysteresis graph')

h3('3.7.6. Polling Interval')
eq('Server utilization = 25ms / 1500ms = 1.7 percent')

fig('figures/figure-3-8-state-machine.png', 'Figure 3-8: System state machine')
fig('figures/figure-3-10-polling-sequence.png', 'Figure 3-10: Polling sequence diagram')
fig('figures/figure-3-11-optimistic-ui-flow.png', 'Figure 3-11: Optimistic UI flowchart')
fig('figures/figure-3-12-communication-flow.png', 'Figure 3-12: Communication flow diagram')

pagebreak()

# ══════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════
h1('4. Implementation')

h2('4.1. Hardware Implementation')
body('The hardware was prototyped on a standard 830-point solderless breadboard. Issues discovered during prototyping included DHT11 signal integrity problems from loose breadboard contacts, PIR false triggers from relay EMI, and relay module current requirements exceeding the shared supply capacity.')
fig('figures/figure-4-3-voltage-current.png', 'Figure 4-3: 7805 output voltage vs load current')
fig('figures/figure-4-4-wifi-provisioning-flow.png', 'Figure 4-4: Wi-Fi provisioning flowchart')

pagebreak()

# ══════════════════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════════════════
h1('5. Results and Evaluation')

h2('5.1. Hardware Performance')
body('The average relay switching response time was 45 ms (+/- 8 ms). PIR detection accuracy was 97 out of 100 tests. DHT11 temperature showed mean absolute error of 1.2 deg C.')
fig('figures/figure-5-2-pir-accuracy.png', 'Figure 5-2: PIR detection accuracy')
fig('figures/figure-5-3-dht11-accuracy.png', 'Figure 5-3: DHT11 sensor accuracy')
fig('figures/figure-5-4-power-consumption.png', 'Figure 5-4: Power consumption analysis')

h2('5.2. Firmware Performance')
tbl(['Endpoint', 'Mean (ms)', 'Median (ms)', '95th %ile (ms)', 'Max (ms)'], [
    ['GET /ping', '12', '10', '18', '45'],
    ['GET /status', '28', '25', '42', '85'],
    ['POST /relay/*', '35', '32', '55', '120'],
    ['POST /mode', '30', '28', '48', '95'],
])
txt('Table 5-1: HTTP response latency (ms)', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
fig('figures/figure-5-5-http-latency.png', 'Figure 5-5: HTTP response latency')
fig('figures/figure-5-6-wifi-stability.png', 'Figure 5-6: Wi-Fi connection stability')

h2('5.3. Mobile Application Performance')
fig('figures/figure-5-7-discovery-time.png', 'Figure 5-7: Device discovery time')
fig('figures/figure-5-8-end-to-end-latency.png', 'Figure 5-8: End-to-end latency breakdown')
fig('figures/figure-5-9-sync-accuracy.png', 'Figure 5-9: State synchronization accuracy')

h2('5.4. System-Level Evaluation')
fig('figures/figure-5-10-7day-operation.png', 'Figure 5-10: 7-day operation timeline')

h3('5.4.3. Cost Analysis')
tbl(['Component', 'Qty', 'Unit (N)', 'Total (N)'], [
    ['ESP32 Dev Board', '1', '6,500', '6,500'],
    ['4-channel relay', '1', '3,000', '3,000'],
    ['DHT11 sensor', '1', '1,500', '1,500'],
    ['PIR sensor', '1', '2,000', '2,000'],
    ['7805 regulator', '1', '500', '500'],
    ['Other components', '-', '-', '4,700'],
    ['Total', '', '', '19,200'],
])
txt('Table 5-2: Bill of materials (Naira)', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
fig('figures/figure-5-11-cost-comparison.png', 'Figure 5-11: Cost comparison')

h2('5.5. Comparison with Existing Solutions')
tbl(['Feature', 'Proposed', 'Alexa', 'Home Asst.', 'Blynk'], [
    ['Cloud', 'None', 'Full', 'Optional', 'Full'],
    ['Latency', '<150ms', '200-500ms', '100-300ms', '300-800ms'],
    ['Discovery', 'Auto', 'Manual', 'Manual', 'Manual'],
    ['Cost', 'N19,200', 'N75k-300k', 'N52k-112k', 'N6,750+'],
    ['Privacy', 'Full local', 'Cloud', 'Configurable', 'Cloud'],
])
txt('Table 5-3: Comparison with existing solutions', italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
fig('figures/figure-5-12-radar-comparison.png', 'Figure 5-12: Multi-axis comparison')

pagebreak()

# ══════════════════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════════════════
h1('6. Concluding Remarks and Future Work')

h2('6.1. Concluding Remarks')
body('This thesis has presented the design, implementation, and evaluation of an ESP32-based IoT home automation system. The research objectives have been achieved:')
for i, c in enumerate([
    'A complete system was implemented at N19,200 total component cost.',
    'The firmware supports zero-configuration Wi-Fi provisioning, automatic device discovery, and a RESTful API.',
    'The mobile application provides real-time control with optimistic UI updates.',
    'Experimental evaluation demonstrated 45ms relay response, 97% PIR accuracy, and 99.6% state sync accuracy.',
], 1):
    txt(f'{i}. {c}', indent=1.27, after=6)

h2('6.2. Future Work')
for i, f in enumerate([
    'MQTT integration for push-based updates.',
    'Voice control integration.',
    'Energy monitoring using current sensors.',
    'Machine learning for occupancy prediction.',
    'Mesh networking using ESP-NOW.',
    'Security hardening with HTTPS.',
    'Mobile app enhancement with push notifications.',
], 1):
    txt(f'{i}. {f}', indent=1.27, after=6)

pagebreak()

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
h1('8. References')
refs = [
    '[1] L. Atzori et al., "The Internet of Things: A survey," Computer Networks, vol. 54, no. 15, pp. 2787-2805, 2010.',
    '[2] "Smart Home Market Report," Fortune Business Insights, 2021.',
    '[3] R. Harper, Inside the Smart Home. London: Springer, 2003.',
    '[4] "Residential Energy Consumption Survey," U.S. EIA, 2020.',
    '[5] W. Kleiminger et al., "Household occupancy monitoring," ACM UbiComp, 2015.',
    '[6] Espressif Systems, "ESP32 Datasheet," v4.4, 2021.',
    '[7] Espressif Systems, "ESP-IDF Programming Guide," 2022.',
    '[8] J. Gubbi et al., "IoT: A vision," Future Generation Computer Systems, vol. 29, no. 7, 2013.',
    '[9] A. Al-Fuqaha et al., "IoT: A survey," IEEE Comm. Surveys, vol. 17, no. 4, 2015.',
    '[10] J. K. Aggarwal, "Human activity analysis: A review," ACM Computing Surveys, vol. 43, 2011.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# Save
doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT) / 1024:.0f} KB')
